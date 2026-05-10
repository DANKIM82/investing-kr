#!/usr/bin/env python3
"""
docx_renderer.py - Word 리서치 노트 생성

JSON context에서 .docx 생성. 한글 폰트 자동 적용.

CLI:
    python infra/docx_renderer.py \\
        --context reports/.tmp/context.json \\
        --output reports/output.docx \\
        --template templates/research_note_kr.docx  # 옵션
"""

import argparse
import json
import os
import sys
from datetime import datetime


def _set_korean_font(run, font_name="Malgun Gothic"):
    """한글 폰트 적용."""
    try:
        from docx.oxml.ns import qn
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
    except Exception:
        pass


def render_docx(context, output_path, template_path=None):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[error] python-docx 미설치. pip install python-docx", file=sys.stderr)
        return False
    
    # 템플릿 사용 또는 새로 생성
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
    
    company = context.get("company", {})
    report = context.get("report", {})
    
    # 표지
    title = doc.add_heading(f"{company.get('name', '')} ({company.get('ticker', '')})", level=0)
    for run in title.runs:
        _set_korean_font(run)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"리서치 노트 — {report.get('date', datetime.now().strftime('%Y-%m-%d'))}")
    run.font.size = Pt(14)
    _set_korean_font(run)
    
    # 투자의견 / 적정가
    if report.get("rating") or report.get("target_price"):
        rating_p = doc.add_paragraph()
        rating_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rating_text = ""
        if report.get("rating"):
            rating_text += f"투자의견: {report['rating']}    "
        if report.get("target_price"):
            rating_text += f"적정가: {report['target_price']:,}"
            if company.get("currency") == "KRW":
                rating_text += "원"
            elif company.get("currency") == "USD":
                rating_text = rating_text.replace(f"{report['target_price']:,}", f"${report['target_price']:,.2f}")
            elif company.get("currency") == "JPY":
                rating_text = rating_text.replace(f"{report['target_price']:,}", f"¥{report['target_price']:,.0f}")
        run = rating_p.add_run(rating_text)
        run.font.size = Pt(12)
        run.bold = True
        _set_korean_font(run)
    
    doc.add_paragraph()
    
    # Executive Summary
    if context.get("executive_summary"):
        h = doc.add_heading("Executive Summary", level=1)
        for run in h.runs:
            _set_korean_font(run)
        p = doc.add_paragraph(context["executive_summary"])
        for run in p.runs:
            _set_korean_font(run)
    
    # Sections
    for section in context.get("sections", []):
        h = doc.add_heading(section.get("title", ""), level=1)
        for run in h.runs:
            _set_korean_font(run)
        
        if section.get("content"):
            p = doc.add_paragraph(section["content"])
            for run in p.runs:
                _set_korean_font(run)
        
        # 표 추가
        for table_data in section.get("tables", []):
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            if not headers or not rows:
                continue
            t = doc.add_table(rows=1 + len(rows), cols=len(headers))
            t.style = "Light Grid Accent 1"
            
            # 헤더
            for j, header in enumerate(headers):
                cell = t.rows[0].cells[j]
                cell.text = str(header)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    _set_korean_font(run)
            
            # 데이터
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    cell = t.rows[i + 1].cells[j]
                    cell.text = str(val)
                    for run in cell.paragraphs[0].runs:
                        _set_korean_font(run)
            doc.add_paragraph()
        
        # 이미지/차트 삽입
        for img_path in section.get("images", []):
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph()
    
    # 푸터
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firm = report.get("firm_name", "Personal Research")
    run = footer_p.add_run(f"Prepared by {firm}")
    run.italic = True
    run.font.size = Pt(10)
    _set_korean_font(run)
    
    src_p = doc.add_paragraph()
    src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = src_p.add_run("Data sources: DART (KR) / yfinance + SEC EDGAR (US) / yfinance (JP)")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _set_korean_font(run)
    
    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc_p.add_run("This report is for educational purposes only. Not investment advice.")
    run.font.size = Pt(8)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    _set_korean_font(run)
    
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--context", required=True, help="JSON context 파일 경로")
    parser.add_argument("--output", required=True, help="출력 .docx 경로")
    parser.add_argument("--template", help="템플릿 .docx (선택)")
    args = parser.parse_args()
    
    with open(args.context, encoding="utf-8") as f:
        context = json.load(f)
    
    success = render_docx(context, args.output, args.template)
    
    if success:
        print(json.dumps({"output": args.output, "status": "success"}, ensure_ascii=False))
    else:
        sys.exit(1)


if __name__ == "__main__":
    main()
